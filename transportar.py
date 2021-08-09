import pyautogui
import clipboard
import pyperclip
import winsound

#COPIA DADOS DO ARQUIVO

from docx import Document

def alimentarOracleMovimentacao():

    input = Document('processos.docx')

    paragraphs = []

    for para in input.paragraphs:
        p = para.text
        if p != "":
            paragraphs.append(p)

    pyautogui.moveTo(136,367,0.2)
    pyautogui.click()

    for para in input.paragraphs:
        strText = para.text
        print(para.text)
        pyperclip.copy(strText)
        pyautogui.hotkey('ctrl','v')
        pyautogui.press('enter')
        pyautogui.press('enter')
        pyautogui.press('enter')
        pyautogui.press('enter')
        #pyperclip.paste
        pyautogui.PAUSE = 2
        
    print(str(len(paragraphs)) + " processos")

    #winsound.MessageBeep()

def main():
    alimentarOracleMovimentacao()

if __name__ == '__main__':
    main()
